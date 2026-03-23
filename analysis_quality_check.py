#!/usr/bin/env python3
"""
SARA Output Quality Analysis Script
====================================
Parses the DOCX output and DSD source, compares all financial data,
identifies errors in totals/subtotals, and produces a detailed report.
"""

import sys
import os
import re
from collections import defaultdict

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from docx import Document
from skills.parse_dsd import parse_dsd
from utils.number_format import parse_korean_number, format_english_number
from ir_schema import ElementType, StatementType

# ─── Config ───
DOCX_PATH = "backend/outputs/ac368833-2106-42f6-b32a-e5edeafcf53a/SARA_result.docx"
DSD_PATH = "files/Hybe 2025 Eng Report.dsd"


# ═══════════════════════════════════════════
# 1. DOCX Parsing
# ═══════════════════════════════════════════

def parse_docx_table(table):
    """Extract all cell text from a python-docx table into a 2D list."""
    rows = []
    for row in table.rows:
        cells = []
        for cell in row.cells:
            cells.append(cell.text.strip())
        rows.append(cells)
    return rows


def parse_number(text):
    """Parse a number from English-formatted financial text."""
    if not text:
        return None
    cleaned = text.strip()
    # Remove backslash (won symbol proxy used in English FS)
    cleaned = cleaned.replace("\\", "").replace("\u20a9", "").strip()
    # Handle dash = zero
    if cleaned in ("-", "—", "–", ""):
        return 0
    # Percentage — skip
    if "%" in cleaned:
        return None
    # Remove text-only
    if re.match(r'^[A-Za-z\s]+$', cleaned):
        return None

    is_negative = False
    if cleaned.startswith("(") and cleaned.endswith(")"):
        is_negative = True
        cleaned = cleaned[1:-1].strip()
    elif cleaned.startswith("-"):
        # Could be a negative number
        rest = cleaned[1:].strip().replace(",", "")
        if rest and rest.replace(".", "").isdigit():
            is_negative = True
            cleaned = rest

    cleaned = cleaned.replace(",", "")
    if not cleaned:
        return None
    try:
        if "." in cleaned:
            val = float(cleaned)
        else:
            val = int(cleaned)
        return -val if is_negative else val
    except ValueError:
        return None


def identify_docx_fs_tables(doc):
    """
    Identify the 4 main financial statement tables in DOCX.
    BS, IS, CE, CF typically appear early and are large tables with
    'in thousands of Korean won' header.
    """
    fs_tables = {}
    # Heuristic: look at early tables for FS patterns
    for i, table in enumerate(doc.tables):
        rows = parse_docx_table(table)
        if not rows:
            continue
        # Combine first 3 rows of text for detection
        sample = " ".join(" ".join(r) for r in rows[:5]).lower()

        # BS detection
        if ("current assets" in sample or "non-current assets" in sample or
            "total assets" in sample) and "BS" not in fs_tables:
            fs_tables["BS"] = (i, rows)
        # IS detection
        elif ("revenue" in sample or "operating profit" in sample or
              "profit for the" in sample or "net income" in sample) and "IS" not in fs_tables:
            # Distinguish from CE
            if "retained earnings" not in sample and "share capital" not in sample:
                fs_tables["IS"] = (i, rows)
        # CE detection
        elif ("share capital" in sample or "capital surplus" in sample or
              "retained earnings" in sample or "equity" in sample.split("\n")[0] if sample else False) and "CE" not in fs_tables:
            # CE has complex multi-column header with equity components
            header_text = " ".join(" ".join(r) for r in rows[:3]).lower()
            if "share capital" in header_text or "retained earnings" in header_text:
                fs_tables["CE"] = (i, rows)
        # CF detection
        elif ("operating activities" in sample or "investing activities" in sample or
              "financing activities" in sample or "cash flows" in sample) and "CF" not in fs_tables:
            fs_tables["CF"] = (i, rows)

    return fs_tables


# ═══════════════════════════════════════════
# 2. DSD → Expected values
# ═══════════════════════════════════════════

def get_dsd_fs_data(parsed_dsd):
    """Extract financial statement data from parsed DSD."""
    fs_data = {}
    for section in parsed_dsd.sections:
        for fs in section.financial_statements:
            st = fs.statement_type.value
            table = fs.table
            if table is None:
                continue
            rows_data = []
            for row in table.headers:
                row_cells = [c.text for c in row.cells]
                rows_data.append(("header", row_cells))
            for row in table.rows:
                row_cells = [c.text for c in row.cells]
                flags = []
                if row.is_total:
                    flags.append("TOTAL")
                if row.is_subtotal:
                    flags.append("SUBTOTAL")
                rows_data.append(("data", row_cells, flags))
            fs_data[st] = rows_data
    return fs_data


def get_dsd_note_tables(parsed_dsd):
    """Extract all note tables from parsed DSD."""
    note_tables = {}
    for section in parsed_dsd.sections:
        for note in section.notes:
            table_idx = 0
            for elem in note.elements:
                if elem.type == ElementType.TABLE and elem.table:
                    key = f"Note {note.number}: {note.title} (Table {table_idx})"
                    table_data = []
                    for row in elem.table.headers:
                        table_data.append([c.text for c in row.cells])
                    for row in elem.table.rows:
                        table_data.append([c.text for c in row.cells])
                    note_tables[key] = table_data
                    table_idx += 1
    return note_tables


# ═══════════════════════════════════════════
# 3. Comparison Logic
# ═══════════════════════════════════════════

def compare_fs_numbers(dsd_rows, docx_rows, fs_type):
    """
    Compare DSD (Korean) numbers to DOCX (English) numbers for a financial statement.
    Returns detailed comparison results.
    """
    results = []

    # For BS/IS/CF: structure is label | current | prior (simplified)
    # DSD uses Korean labels, DOCX uses English labels
    # We match by position (row index) and compare numeric columns

    # Extract numeric values from DSD rows
    dsd_numeric = []
    for item in dsd_rows:
        if item[0] == "header":
            continue
        row_cells = item[1]
        flags = item[2] if len(item) > 2 else []
        # Find label (first non-empty cell)
        label = ""
        numbers = []
        for cell in row_cells:
            val = parse_korean_number(cell)
            if val is not None and cell.strip() not in ("", "-"):
                numbers.append((cell.strip(), val))
            elif not label and cell.strip():
                label = cell.strip()
        dsd_numeric.append({
            "label": label,
            "numbers": numbers,
            "flags": flags,
            "raw": row_cells,
        })

    # Extract from DOCX
    docx_numeric = []
    for row_cells in docx_rows:
        label = ""
        numbers = []
        for cell in row_cells:
            val = parse_number(cell)
            if val is not None and cell.strip() not in ("",):
                numbers.append((cell.strip(), val))
            elif not label and cell.strip() and not re.match(r'^[\d,\(\)\-\.\\]+$', cell.strip()):
                label = cell.strip()
        docx_numeric.append({
            "label": label,
            "numbers": numbers,
            "raw": row_cells,
        })

    return dsd_numeric, docx_numeric


def find_total_rows_dsd(dsd_rows):
    """Find all total/subtotal rows in DSD data."""
    totals = []
    for item in dsd_rows:
        if item[0] == "header":
            continue
        row_cells = item[1]
        flags = item[2] if len(item) > 2 else []
        combined = " ".join(row_cells)
        if "TOTAL" in flags or "SUBTOTAL" in flags or re.search(r'합\s*계|총\s*계|소\s*계', combined):
            label = ""
            numbers = []
            for cell in row_cells:
                val = parse_korean_number(cell)
                if val is not None:
                    numbers.append(val)
                elif not label and cell.strip():
                    label = cell.strip()
            totals.append({"label": label, "values": numbers, "flags": flags})
    return totals


def find_total_rows_docx(docx_rows):
    """Find all total/subtotal rows in DOCX data."""
    totals = []
    for row_cells in docx_rows:
        combined = " ".join(row_cells).lower()
        if "total" in combined or "subtotal" in combined or "sub-total" in combined:
            label = ""
            numbers = []
            for cell in row_cells:
                val = parse_number(cell)
                if val is not None:
                    numbers.append(val)
                elif not label and cell.strip():
                    label = cell.strip()
            totals.append({"label": label, "values": numbers})
    return totals


# ═══════════════════════════════════════════
# 4. Main Analysis
# ═══════════════════════════════════════════

def main():
    print("=" * 100)
    print("SARA OUTPUT QUALITY ANALYSIS REPORT")
    print("=" * 100)
    print()

    # Load files
    print("Loading DOCX output...")
    doc = Document(DOCX_PATH)
    print(f"  Total tables in DOCX: {len(doc.tables)}")

    print("Loading DSD source...")
    parsed_dsd = parse_dsd(DSD_PATH)
    print(f"  Company: {parsed_dsd.meta.company}")
    print(f"  Period: {parsed_dsd.meta.period_current} (current), {parsed_dsd.meta.period_prior} (prior)")
    print()

    # ──────────────────────────────────────────
    # Part A: Financial Statement Tables (BS, IS, CE, CF)
    # ──────────────────────────────────────────
    print("=" * 100)
    print("PART A: MAIN FINANCIAL STATEMENTS (BS, IS, CE, CF)")
    print("=" * 100)

    dsd_fs = get_dsd_fs_data(parsed_dsd)
    docx_fs = identify_docx_fs_tables(doc)

    # We need to find which DOCX tables correspond to BS/IS/CE/CF
    # Let's also do a broader search
    print("\n--- Scanning all DOCX tables for FS patterns ---")
    for i, table in enumerate(doc.tables):
        rows = parse_docx_table(table)
        if len(rows) < 5:
            continue
        first_row_text = " ".join(rows[0]).lower()
        second_row_text = " ".join(rows[1]).lower() if len(rows) > 1 else ""
        combined = first_row_text + " " + second_row_text

        # Check for FS indicators
        if any(kw in combined for kw in ["total assets", "current assets", "non-current assets",
                                          "total liabilities", "total equity"]):
            print(f"  Table {i} ({len(rows)}r): Likely BS — {rows[1][:3] if len(rows)>1 else 'N/A'}")
        elif any(kw in combined for kw in ["revenue", "operating profit", "profit for the year",
                                            "earnings per share"]):
            print(f"  Table {i} ({len(rows)}r): Likely IS — first cell: {rows[0][0][:40]}")
        elif any(kw in combined for kw in ["share capital", "capital surplus", "retained earnings",
                                            "treasury shares"]):
            print(f"  Table {i} ({len(rows)}r): Likely CE — first cell: {rows[0][0][:40]}")
        elif any(kw in combined for kw in ["operating activities", "investing activities",
                                            "financing activities", "cash and cash equivalents"]):
            print(f"  Table {i} ({len(rows)}r): Likely CF — first cell: {rows[0][0][:40]}")

    # Find the actual FS tables more carefully
    # For Hybe, the FS tables should be among the first large tables
    # Let's manually find them
    print("\n--- Looking at large tables (>20 rows) ---")
    large_tables = []
    for i, table in enumerate(doc.tables):
        rows = parse_docx_table(table)
        if len(rows) >= 15:
            sample = " ".join(" ".join(r) for r in rows[:3])[:120]
            print(f"  Table {i}: {len(rows)} rows — {sample}")
            large_tables.append(i)

    # ──────────────────────────────────────────
    # Detailed BS analysis
    # ──────────────────────────────────────────
    fs_type_map = {"BS": "재무상태표", "IS": "포괄손익계산서", "CE": "자본변동표", "CF": "현금흐름표"}

    # Try to find the DOCX FS tables manually by scanning for known patterns
    # The DOCX typically has the 4 FS right after the cover page tables
    docx_fs_idx = {}
    for i, table in enumerate(doc.tables):
        rows = parse_docx_table(table)
        if len(rows) < 10:
            continue
        all_text = " ".join(" ".join(r) for r in rows)

        if "Total assets" in all_text and "Total liabilities" in all_text and "BS" not in docx_fs_idx:
            docx_fs_idx["BS"] = i
        elif ("Revenue" in all_text or "Operating profit" in all_text) and \
             "Profit for the year" in all_text and "IS" not in docx_fs_idx and \
             "Share capital" not in " ".join(rows[0]):
            docx_fs_idx["IS"] = i
        elif "Cash flows from operating" in all_text and "CF" not in docx_fs_idx:
            docx_fs_idx["CF"] = i

    # CE is special — look for it
    for i, table in enumerate(doc.tables):
        rows = parse_docx_table(table)
        if len(rows) < 8:
            continue
        header_text = " ".join(" ".join(r) for r in rows[:3])
        if ("Share capital" in header_text or "Capital surplus" in header_text) and \
           ("Retained earnings" in header_text or "Treasury" in header_text):
            docx_fs_idx["CE"] = i
            break

    print(f"\n--- Identified FS tables in DOCX: {docx_fs_idx}")
    print()

    total_correct = 0
    total_wrong = 0
    total_missing = 0
    total_extra = 0
    error_details = []

    for fs_code in ["BS", "IS", "CE", "CF"]:
        print(f"\n{'─' * 80}")
        print(f"  {fs_code} ({fs_type_map[fs_code]})")
        print(f"{'─' * 80}")

        if fs_code not in dsd_fs:
            print(f"  ⚠ DSD has no {fs_code} data!")
            continue
        if fs_code not in docx_fs_idx:
            print(f"  ⚠ Could not find {fs_code} table in DOCX!")
            continue

        dsd_rows = dsd_fs[fs_code]
        docx_table_idx = docx_fs_idx[fs_code]
        docx_rows = parse_docx_table(doc.tables[docx_table_idx])

        # Skip CE for now since it has a very different structure
        if fs_code == "CE":
            print(f"  DOCX Table {docx_table_idx}: {len(docx_rows)} rows")
            print(f"  DSD: {len(dsd_rows)} rows (incl headers)")
            print(f"  [CE has complex multi-column structure — analyzed separately below]")
            continue

        print(f"  DOCX Table {docx_table_idx}: {len(docx_rows)} rows")
        print(f"  DSD: {len(dsd_rows)} rows (incl headers)")

        # For BS/IS/CF: compare row by row
        # DSD structure: label | 당기 numbers | 전기 numbers
        # DOCX structure: label | current year | prior year

        # Extract DSD numeric data (skip headers)
        dsd_data = []
        for item in dsd_rows:
            if item[0] == "header":
                continue
            cells = item[1]
            flags = item[2] if len(item) > 2 else []
            # Find label and numbers
            label = ""
            nums = []
            for c in cells:
                txt = c.strip()
                if not txt:
                    continue
                val = parse_korean_number(txt)
                if val is not None:
                    nums.append(val)
                elif not label:
                    label = txt
            dsd_data.append({"label": label, "nums": nums, "flags": flags})

        # Extract DOCX numeric data
        docx_data = []
        for row_cells in docx_rows:
            label = ""
            nums = []
            for c in row_cells:
                txt = c.strip()
                if not txt:
                    continue
                val = parse_number(txt)
                if val is not None:
                    nums.append(val)
                elif not label and not re.match(r'^[\d,\(\)\-\.\\]+$', txt):
                    # Check it's not just "in thousands..."
                    if "in thousands" in txt.lower() or "in shares" in txt.lower():
                        continue
                    label = txt
            docx_data.append({"label": label, "nums": nums})

        # Compare: match by row position (after filtering out header-only rows)
        # Filter out rows with no numbers and no label
        dsd_filtered = [d for d in dsd_data if d["label"] or d["nums"]]
        docx_filtered = [d for d in docx_data if d["label"] or d["nums"]]

        # Remove pure header rows from DOCX (first 1-2 rows)
        # Detect header rows: they typically have year strings like "2025", "2024"
        docx_filtered_clean = []
        for d in docx_filtered:
            if d["label"] and any(y in d["label"] for y in ["2025", "2024", "2023", "thousands"]):
                continue
            if d["nums"] == [2025, 2024] or d["nums"] == [2024, 2023]:
                continue
            docx_filtered_clean.append(d)
        docx_filtered = docx_filtered_clean

        print(f"\n  DSD data rows: {len(dsd_filtered)}")
        print(f"  DOCX data rows: {len(docx_filtered)}")

        # Now match by position and compare numbers
        max_rows = max(len(dsd_filtered), len(docx_filtered))
        stmt_correct = 0
        stmt_wrong = 0
        stmt_missing = 0

        print(f"\n  {'Row':>4} | {'DSD Label':40} | {'DSD Nums':>30} | {'DOCX Nums':>30} | Status")
        print(f"  {'─'*4}─┼─{'─'*40}─┼─{'─'*30}─┼─{'─'*30}─┼─{'─'*15}")

        for idx in range(max_rows):
            dsd_row = dsd_filtered[idx] if idx < len(dsd_filtered) else None
            docx_row = docx_filtered[idx] if idx < len(docx_filtered) else None

            if dsd_row is None:
                total_extra += 1
                continue
            if docx_row is None:
                total_missing += 1
                stmt_missing += 1
                continue

            dsd_label = dsd_row["label"][:40]
            dsd_nums = dsd_row["nums"]
            docx_nums = docx_row["nums"]

            # Deduplicate consecutive identical values (from merged cells in DOCX)
            def dedup(nums):
                if len(nums) <= 2:
                    return nums
                result = [nums[0]]
                for n in nums[1:]:
                    if n != result[-1]:
                        result.append(n)
                return result

            docx_nums_dedup = dedup(docx_nums)

            # Compare: DSD may have more columns (당기 + 전기 with sub-columns)
            # The main comparison is: do the numeric values match?
            dsd_str = ", ".join(f"{n:,}" for n in dsd_nums[:4])
            docx_str = ", ".join(f"{n:,}" for n in docx_nums_dedup[:4])

            # Check match
            match = False
            if len(dsd_nums) >= 2 and len(docx_nums_dedup) >= 2:
                # For BS: typically 2 numbers (current, prior)
                if dsd_nums[:2] == docx_nums_dedup[:2]:
                    match = True
            elif len(dsd_nums) == 1 and len(docx_nums_dedup) >= 1:
                if dsd_nums[0] == docx_nums_dedup[0]:
                    match = True
            elif len(dsd_nums) == 0 and len(docx_nums_dedup) == 0:
                match = True  # Both have no numbers (label-only rows)

            is_total = "TOTAL" in dsd_row.get("flags", []) or "SUBTOTAL" in dsd_row.get("flags", [])
            flag_str = " [TOTAL]" if is_total else ""

            if match:
                status = "OK"
                stmt_correct += 1
                total_correct += 1
            else:
                status = "MISMATCH"
                stmt_wrong += 1
                total_wrong += 1
                error_details.append({
                    "fs": fs_code,
                    "row": idx,
                    "label": dsd_label,
                    "expected": dsd_nums[:4],
                    "actual": docx_nums_dedup[:4],
                    "is_total": is_total,
                })

            # Only print mismatches and totals
            if not match or is_total:
                print(f"  {idx:4} | {dsd_label:40} | {dsd_str:>30} | {docx_str:>30} | {status}{flag_str}")

        print(f"\n  Summary: {stmt_correct} correct, {stmt_wrong} wrong, {stmt_missing} missing")

    # ──────────────────────────────────────────
    # Part B: Total/Subtotal Analysis
    # ──────────────────────────────────────────
    print("\n")
    print("=" * 100)
    print("PART B: TOTAL / SUBTOTAL ANALYSIS")
    print("=" * 100)

    for fs_code in ["BS", "IS", "CF"]:
        if fs_code not in dsd_fs or fs_code not in docx_fs_idx:
            continue

        print(f"\n--- {fs_code} Totals ---")

        # DSD totals
        dsd_totals = find_total_rows_dsd(dsd_fs[fs_code])
        docx_totals = find_total_rows_docx(parse_docx_table(doc.tables[docx_fs_idx[fs_code]]))

        print(f"  DSD total rows: {len(dsd_totals)}")
        for t in dsd_totals:
            vals = ", ".join(f"{v:,}" for v in t["values"][:4])
            print(f"    {t['label'][:50]:50} | {vals}")

        print(f"  DOCX total rows: {len(docx_totals)}")
        for t in docx_totals:
            vals = ", ".join(f"{v:,}" for v in t["values"][:4])
            print(f"    {t['label'][:50]:50} | {vals}")

        # Cross-compare totals
        print(f"\n  Cross-comparison:")
        for di, dt in enumerate(dsd_totals):
            matched = False
            for doi, dot in enumerate(docx_totals):
                # Compare values
                dsd_vals = dt["values"][:2] if dt["values"] else []
                docx_vals = dot["values"][:2] if dot["values"] else []
                if dsd_vals and docx_vals and dsd_vals == docx_vals:
                    print(f"    OK: DSD[{dt['label'][:30]}] == DOCX[{dot['label'][:30]}] = {dsd_vals}")
                    matched = True
                    break
            if not matched and dt["values"]:
                dsd_vals_str = ", ".join(f"{v:,}" for v in dt["values"][:4])
                # Find closest DOCX total by label similarity
                closest_docx = ""
                for dot in docx_totals:
                    docx_vals_str = ", ".join(f"{v:,}" for v in dot["values"][:4])
                    closest_docx = f"DOCX closest: {dot['label'][:30]} = {docx_vals_str}"
                    break
                print(f"    MISMATCH: DSD[{dt['label'][:30]}] = {dsd_vals_str}")
                if closest_docx:
                    print(f"              {closest_docx}")

    # ──────────────────────────────────────────
    # Part C: Note Table Analysis
    # ──────────────────────────────────────────
    print("\n")
    print("=" * 100)
    print("PART C: NOTE TABLE ANALYSIS")
    print("=" * 100)

    dsd_note_tables = get_dsd_note_tables(parsed_dsd)
    print(f"\nTotal DSD note tables: {len(dsd_note_tables)}")

    # For each DOCX table (after the FS tables), check if it has data
    note_analysis = {
        "populated": 0,
        "empty_or_prior_only": 0,
        "partial": 0,
    }

    fs_table_indices = set(docx_fs_idx.values())

    # Check each DOCX table for emptiness
    print(f"\n--- DOCX Note Tables Status ---")
    print(f"{'Idx':>4} | {'Rows':>4} | {'Status':15} | First Row Preview")
    print(f"{'─'*4}─┼─{'─'*4}─┼─{'─'*15}─┼─{'─'*60}")

    empty_tables = []
    populated_tables = []
    partial_tables = []

    for i, table in enumerate(doc.tables):
        if i in fs_table_indices:
            continue  # Skip FS tables
        rows = parse_docx_table(table)
        if len(rows) < 2:
            continue

        # Count cells with numeric data
        total_cells = 0
        filled_cells = 0
        zero_cells = 0
        for r_idx, row in enumerate(rows):
            if r_idx == 0:
                continue  # skip header
            for cell in row:
                if cell.strip() and cell.strip() not in ("", "(in thousands of Korean won)", "(in shares)"):
                    val = parse_number(cell)
                    if val is not None:
                        total_cells += 1
                        if val == 0 and cell.strip() in ("-", "—", "–"):
                            zero_cells += 1
                        filled_cells += 1

        first_row = " | ".join(c[:20] for c in rows[0][:4])

        if total_cells == 0:
            status = "EMPTY"
            empty_tables.append(i)
        elif filled_cells < total_cells * 0.3:
            status = "SPARSE"
            partial_tables.append(i)
        else:
            status = "POPULATED"
            populated_tables.append(i)

        # Only report non-populated
        if status != "POPULATED":
            print(f"  {i:4} | {len(rows):4} | {status:15} | {first_row[:60]}")

    print(f"\n  Summary: {len(populated_tables)} populated, {len(partial_tables)} sparse, {len(empty_tables)} empty")

    # ──────────────────────────────────────────
    # Part D: Sample cell-by-cell comparison for key note tables
    # ──────────────────────────────────────────
    print("\n")
    print("=" * 100)
    print("PART D: SAMPLE CELL-BY-CELL NOTE TABLE COMPARISONS")
    print("=" * 100)

    # Match DSD note tables to DOCX tables by content
    # Strategy: for each DSD note table, find the closest DOCX table
    dsd_notes_list = []
    for section in parsed_dsd.sections:
        for note in section.notes:
            for elem in note.elements:
                if elem.type == ElementType.TABLE and elem.table:
                    rows = []
                    for r in elem.table.headers:
                        rows.append([c.text for c in r.cells])
                    for r in elem.table.rows:
                        rows.append([c.text for c in r.cells])
                    # Get all numbers
                    numbers = []
                    for row in rows:
                        for cell in row:
                            val = parse_korean_number(cell)
                            if val is not None and val != 0:
                                numbers.append(val)
                    dsd_notes_list.append({
                        "note": f"{note.number}. {note.title}",
                        "rows": rows,
                        "numbers": numbers,
                        "row_count": len(rows),
                    })

    print(f"\nDSD note tables: {len(dsd_notes_list)}")

    # For the first ~10 DSD note tables with numbers, try to find matching DOCX table
    checked = 0
    matches_found = 0
    mismatches_found = 0

    for dsd_nt in dsd_notes_list:
        if not dsd_nt["numbers"] or checked >= 20:
            continue
        checked += 1

        # Try to find matching DOCX table
        best_match = None
        best_score = 0

        for di, table in enumerate(doc.tables):
            if di in fs_table_indices:
                continue
            docx_rows = parse_docx_table(table)

            # Compare row counts (approximate match)
            if abs(len(docx_rows) - dsd_nt["row_count"]) > 5:
                continue

            # Compare numbers
            docx_numbers = []
            for row in docx_rows:
                for cell in row:
                    val = parse_number(cell)
                    if val is not None and val != 0:
                        docx_numbers.append(val)

            # Score: count matching numbers
            if not docx_numbers:
                continue
            matching = 0
            for dnum in dsd_nt["numbers"][:10]:
                if dnum in docx_numbers:
                    matching += 1
            score = matching / max(len(dsd_nt["numbers"][:10]), 1)
            if score > best_score:
                best_score = score
                best_match = (di, docx_rows, docx_numbers, score)

        print(f"\n  DSD: {dsd_nt['note'][:60]} ({dsd_nt['row_count']} rows, {len(dsd_nt['numbers'])} numbers)")
        if best_match:
            di, docx_rows, docx_numbers, score = best_match
            print(f"  DOCX Table {di} (match score: {score:.0%})")

            # Detailed comparison: check which DSD numbers appear in DOCX
            dsd_nums_set = set(dsd_nt["numbers"])
            docx_nums_set = set(docx_numbers)
            common = dsd_nums_set & docx_nums_set
            only_dsd = dsd_nums_set - docx_nums_set
            only_docx = docx_nums_set - dsd_nums_set

            if only_dsd:
                matches_found += 1 if common else 0
                mismatches_found += 1
                print(f"    Common: {len(common)}, Only in DSD: {len(only_dsd)}, Only in DOCX: {len(only_docx)}")
                if only_dsd:
                    sample = sorted(only_dsd, key=abs, reverse=True)[:5]
                    print(f"    Missing from DOCX (top 5 by magnitude): {[f'{v:,}' for v in sample]}")
            else:
                matches_found += 1
                print(f"    All {len(common)} numbers match!")
        else:
            print(f"    No matching DOCX table found!")
            mismatches_found += 1

    # ──────────────────────────────────────────
    # Part E: Row-by-row comparison of first few FS tables
    # ──────────────────────────────────────────
    print("\n")
    print("=" * 100)
    print("PART E: DETAILED ROW-BY-ROW BS COMPARISON")
    print("=" * 100)

    if "BS" in dsd_fs and "BS" in docx_fs_idx:
        dsd_rows = dsd_fs["BS"]
        docx_rows = parse_docx_table(doc.tables[docx_fs_idx["BS"]])

        print(f"\n  DSD BS: {len(dsd_rows)} entries, DOCX BS: {len(docx_rows)} rows")
        print()
        print(f"  {'#':>3} | {'DSD Label':35} | {'DSD Current':>18} | {'DSD Prior':>18} | {'DOCX Current':>18} | {'DOCX Prior':>18} | Match")
        print(f"  {'─'*3}─┼─{'─'*35}─┼─{'─'*18}─┼─{'─'*18}─┼─{'─'*18}─┼─{'─'*18}─┼─{'─'*6}")

        dsd_data_rows = []
        for item in dsd_rows:
            if item[0] == "header":
                continue
            cells = item[1]
            flags = item[2] if len(item) > 2 else []
            label = ""
            nums = []
            for c in cells:
                txt = c.strip()
                if not txt:
                    continue
                val = parse_korean_number(txt)
                if val is not None:
                    nums.append(val)
                elif not label:
                    label = txt
            dsd_data_rows.append({"label": label, "nums": nums, "flags": flags})

        # DOCX data (skip header rows)
        docx_data_rows = []
        for row in docx_rows:
            label = ""
            nums = []
            for c in row:
                txt = c.strip()
                if not txt:
                    continue
                val = parse_number(txt)
                if val is not None:
                    nums.append(val)
                elif not label and "thousands" not in txt.lower() and "shares" not in txt.lower():
                    if not re.match(r'^[\d,]+$', txt):
                        label = txt
            docx_data_rows.append({"label": label, "nums": nums})

        # Remove header rows from DOCX
        docx_data_rows = [d for d in docx_data_rows
                          if d["label"] not in ("", "2025", "2024") and
                          d["nums"] != [2025, 2024]]

        def dedup(nums):
            if len(nums) <= 2:
                return nums
            result = [nums[0]]
            for n in nums[1:]:
                if n != result[-1]:
                    result.append(n)
            return result

        row_idx = 0
        for dsd_r in dsd_data_rows:
            if not dsd_r["label"] and not dsd_r["nums"]:
                continue
            # Find next non-empty DOCX row
            docx_r = None
            while row_idx < len(docx_data_rows):
                if docx_data_rows[row_idx]["label"] or docx_data_rows[row_idx]["nums"]:
                    docx_r = docx_data_rows[row_idx]
                    row_idx += 1
                    break
                row_idx += 1

            dsd_cur = f"{dsd_r['nums'][0]:>15,}" if len(dsd_r['nums']) > 0 else ""
            dsd_pri = f"{dsd_r['nums'][1]:>15,}" if len(dsd_r['nums']) > 1 else ""

            if docx_r:
                dnums = dedup(docx_r["nums"])
                docx_cur = f"{dnums[0]:>15,}" if len(dnums) > 0 else ""
                docx_pri = f"{dnums[1]:>15,}" if len(dnums) > 1 else ""
            else:
                docx_cur = "N/A"
                docx_pri = "N/A"

            # Match check
            match_str = ""
            if docx_r and dsd_r["nums"] and docx_r["nums"]:
                dnums = dedup(docx_r["nums"])
                if dsd_r["nums"][:2] == dnums[:2]:
                    match_str = "OK"
                else:
                    match_str = "WRONG"
            elif not dsd_r["nums"] and (not docx_r or not docx_r["nums"]):
                match_str = "OK"
            else:
                match_str = "MISSING"

            flag = " *" if ("TOTAL" in dsd_r.get("flags", []) or "SUBTOTAL" in dsd_r.get("flags", [])) else ""
            print(f"  {row_idx:3} | {dsd_r['label'][:35]:35} | {dsd_cur:>18} | {dsd_pri:>18} | {docx_cur:>18} | {docx_pri:>18} | {match_str}{flag}")


    # ──────────────────────────────────────────
    # Part F: IS Detailed
    # ──────────────────────────────────────────
    print("\n")
    print("=" * 100)
    print("PART F: DETAILED ROW-BY-ROW IS COMPARISON")
    print("=" * 100)

    if "IS" in dsd_fs and "IS" in docx_fs_idx:
        dsd_rows = dsd_fs["IS"]
        docx_rows = parse_docx_table(doc.tables[docx_fs_idx["IS"]])

        dsd_data_rows = []
        for item in dsd_rows:
            if item[0] == "header":
                continue
            cells = item[1]
            flags = item[2] if len(item) > 2 else []
            label = ""
            nums = []
            for c in cells:
                txt = c.strip()
                if not txt:
                    continue
                val = parse_korean_number(txt)
                if val is not None:
                    nums.append(val)
                elif not label:
                    label = txt
            dsd_data_rows.append({"label": label, "nums": nums, "flags": flags})

        docx_data_rows = []
        for row in docx_rows:
            label = ""
            nums = []
            for c in row:
                txt = c.strip()
                if not txt:
                    continue
                val = parse_number(txt)
                if val is not None:
                    nums.append(val)
                elif not label and "thousands" not in txt.lower():
                    if not re.match(r'^[\d,]+$', txt):
                        label = txt
            docx_data_rows.append({"label": label, "nums": nums})

        docx_data_rows = [d for d in docx_data_rows
                          if d["label"] not in ("", "2025", "2024") and
                          d["nums"] != [2025, 2024]]

        def dedup(nums):
            if len(nums) <= 2:
                return nums
            result = [nums[0]]
            for n in nums[1:]:
                if n != result[-1]:
                    result.append(n)
            return result

        print(f"\n  {'#':>3} | {'DSD Label':40} | {'DSD Current':>18} | {'DSD Prior':>18} | {'DOCX Current':>18} | {'DOCX Prior':>18} | Match")
        print(f"  {'─'*3}─┼─{'─'*40}─┼─{'─'*18}─┼─{'─'*18}─┼─{'─'*18}─┼─{'─'*18}─┼─{'─'*6}")

        row_idx = 0
        for dsd_r in dsd_data_rows:
            if not dsd_r["label"] and not dsd_r["nums"]:
                continue
            docx_r = None
            while row_idx < len(docx_data_rows):
                if docx_data_rows[row_idx]["label"] or docx_data_rows[row_idx]["nums"]:
                    docx_r = docx_data_rows[row_idx]
                    row_idx += 1
                    break
                row_idx += 1

            dsd_cur = f"{dsd_r['nums'][0]:>15,}" if len(dsd_r['nums']) > 0 else ""
            dsd_pri = f"{dsd_r['nums'][1]:>15,}" if len(dsd_r['nums']) > 1 else ""

            if docx_r:
                dnums = dedup(docx_r["nums"])
                docx_cur = f"{dnums[0]:>15,}" if len(dnums) > 0 else ""
                docx_pri = f"{dnums[1]:>15,}" if len(dnums) > 1 else ""
            else:
                docx_cur = "N/A"
                docx_pri = "N/A"

            match_str = ""
            if docx_r and dsd_r["nums"] and docx_r["nums"]:
                dnums = dedup(docx_r["nums"])
                if dsd_r["nums"][:2] == dnums[:2]:
                    match_str = "OK"
                else:
                    match_str = "WRONG"
            elif not dsd_r["nums"] and (not docx_r or not docx_r["nums"]):
                match_str = "OK"
            else:
                match_str = "MISSING"

            flag = " *" if ("TOTAL" in dsd_r.get("flags", []) or "SUBTOTAL" in dsd_r.get("flags", [])) else ""
            print(f"  {row_idx:3} | {dsd_r['label'][:40]:40} | {dsd_cur:>18} | {dsd_pri:>18} | {docx_cur:>18} | {docx_pri:>18} | {match_str}{flag}")


    # ──────────────────────────────────────────
    # Part G: CF Detailed
    # ──────────────────────────────────────────
    print("\n")
    print("=" * 100)
    print("PART G: DETAILED ROW-BY-ROW CF COMPARISON")
    print("=" * 100)

    if "CF" in dsd_fs and "CF" in docx_fs_idx:
        dsd_rows = dsd_fs["CF"]
        docx_rows = parse_docx_table(doc.tables[docx_fs_idx["CF"]])

        dsd_data_rows = []
        for item in dsd_rows:
            if item[0] == "header":
                continue
            cells = item[1]
            flags = item[2] if len(item) > 2 else []
            label = ""
            nums = []
            for c in cells:
                txt = c.strip()
                if not txt:
                    continue
                val = parse_korean_number(txt)
                if val is not None:
                    nums.append(val)
                elif not label:
                    label = txt
            dsd_data_rows.append({"label": label, "nums": nums, "flags": flags})

        docx_data_rows = []
        for row in docx_rows:
            label = ""
            nums = []
            for c in row:
                txt = c.strip()
                if not txt:
                    continue
                val = parse_number(txt)
                if val is not None:
                    nums.append(val)
                elif not label and "thousands" not in txt.lower():
                    if not re.match(r'^[\d,]+$', txt):
                        label = txt
            docx_data_rows.append({"label": label, "nums": nums})

        docx_data_rows = [d for d in docx_data_rows
                          if d["label"] not in ("", "2025", "2024") and
                          d["nums"] != [2025, 2024]]

        def dedup(nums):
            if len(nums) <= 2:
                return nums
            result = [nums[0]]
            for n in nums[1:]:
                if n != result[-1]:
                    result.append(n)
            return result

        print(f"\n  {'#':>3} | {'DSD Label':45} | {'DSD Current':>18} | {'DSD Prior':>18} | {'DOCX Current':>18} | {'DOCX Prior':>18} | Match")
        print(f"  {'─'*3}─┼─{'─'*45}─┼─{'─'*18}─┼─{'─'*18}─┼─{'─'*18}─┼─{'─'*18}─┼─{'─'*6}")

        row_idx = 0
        for dsd_r in dsd_data_rows:
            if not dsd_r["label"] and not dsd_r["nums"]:
                continue
            docx_r = None
            while row_idx < len(docx_data_rows):
                if docx_data_rows[row_idx]["label"] or docx_data_rows[row_idx]["nums"]:
                    docx_r = docx_data_rows[row_idx]
                    row_idx += 1
                    break
                row_idx += 1

            dsd_cur = f"{dsd_r['nums'][0]:>15,}" if len(dsd_r['nums']) > 0 else ""
            dsd_pri = f"{dsd_r['nums'][1]:>15,}" if len(dsd_r['nums']) > 1 else ""

            if docx_r:
                dnums = dedup(docx_r["nums"])
                docx_cur = f"{dnums[0]:>15,}" if len(dnums) > 0 else ""
                docx_pri = f"{dnums[1]:>15,}" if len(dnums) > 1 else ""
            else:
                docx_cur = "N/A"
                docx_pri = "N/A"

            match_str = ""
            if docx_r and dsd_r["nums"] and docx_r["nums"]:
                dnums = dedup(docx_r["nums"])
                if dsd_r["nums"][:2] == dnums[:2]:
                    match_str = "OK"
                else:
                    match_str = "WRONG"
            elif not dsd_r["nums"] and (not docx_r or not docx_r["nums"]):
                match_str = "OK"
            else:
                match_str = "MISSING"

            flag = " *" if ("TOTAL" in dsd_r.get("flags", []) or "SUBTOTAL" in dsd_r.get("flags", [])) else ""
            print(f"  {row_idx:3} | {dsd_r['label'][:45]:45} | {dsd_cur:>18} | {dsd_pri:>18} | {docx_cur:>18} | {docx_pri:>18} | {match_str}{flag}")

    # ──────────────────────────────────────────
    # Part H: Error Pattern Analysis
    # ──────────────────────────────────────────
    print("\n")
    print("=" * 100)
    print("PART H: ERROR PATTERN ANALYSIS")
    print("=" * 100)

    if error_details:
        print(f"\nTotal mismatches in FS tables: {len(error_details)}")

        # Categorize errors
        total_errors = [e for e in error_details if e["is_total"]]
        non_total_errors = [e for e in error_details if not e["is_total"]]

        print(f"  Total/Subtotal errors: {len(total_errors)}")
        print(f"  Regular row errors: {len(non_total_errors)}")

        # Analyze patterns
        off_by_sign = 0
        off_by_factor = 0
        completely_different = 0
        prior_in_current = 0

        for e in error_details:
            if not e["expected"] or not e["actual"]:
                completely_different += 1
                continue

            exp = e["expected"]
            act = e["actual"]

            # Check sign error
            if len(exp) > 0 and len(act) > 0 and exp[0] == -act[0]:
                off_by_sign += 1
            # Check if prior year value is in current year position
            elif len(exp) > 1 and len(act) > 0 and exp[1] == act[0]:
                prior_in_current += 1
            # Check factor errors (1000x etc)
            elif len(exp) > 0 and len(act) > 0 and exp[0] != 0 and act[0] != 0:
                ratio = abs(act[0] / exp[0]) if exp[0] != 0 else 0
                if ratio in (1000, 0.001, 1000000, 0.000001):
                    off_by_factor += 1
                else:
                    completely_different += 1
            else:
                completely_different += 1

        print(f"\n  Error patterns:")
        print(f"    Sign errors (positive/negative flipped): {off_by_sign}")
        print(f"    Prior year in current year position: {prior_in_current}")
        print(f"    Factor errors (1000x, etc.): {off_by_factor}")
        print(f"    Completely different values: {completely_different}")

        # Print specific examples
        if total_errors:
            print(f"\n  Specific TOTAL errors:")
            for e in total_errors[:10]:
                exp_str = ", ".join(f"{v:,}" for v in e["expected"])
                act_str = ", ".join(f"{v:,}" for v in e["actual"])
                print(f"    {e['fs']} - {e['label'][:40]}")
                print(f"      Expected (DSD): {exp_str}")
                print(f"      Actual (DOCX):  {act_str}")
    else:
        print("\nNo mismatches found in FS tables!")

    # ──────────────────────────────────────────
    # Part I: Overall Summary
    # ──────────────────────────────────────────
    print("\n")
    print("=" * 100)
    print("OVERALL SUMMARY")
    print("=" * 100)
    print(f"\n  Main FS Tables:")
    print(f"    Correct cells: {total_correct}")
    print(f"    Wrong cells: {total_wrong}")
    print(f"    Missing cells: {total_missing}")
    print(f"    Extra cells: {total_extra}")
    if (total_correct + total_wrong) > 0:
        accuracy = total_correct / (total_correct + total_wrong) * 100
        print(f"    Accuracy: {accuracy:.1f}%")

    print(f"\n  Note Tables:")
    print(f"    Populated: {len(populated_tables)}")
    print(f"    Sparse: {len(partial_tables)}")
    print(f"    Empty: {len(empty_tables)}")

    print(f"\n  Agent Log Summary:")
    print(f"    Year rolling: 131 elements modified")
    print(f"    Auto-fill: 104 tables, 1,443 cells updated")
    print(f"    DSD tables: 138 total, 104 matched (34 unmatched)")
    print(f"    Glossary: 441 Korean-English pairs built")
    print(f"    Agent steps: 18 steps (08:54:16 to 09:06:06 = ~12 minutes)")


if __name__ == "__main__":
    main()
