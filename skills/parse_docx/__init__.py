"""
parse_docx skill — Parse English financial statement DOCX files into IR.

Entry point: parse_docx(file_path) → ParsedDocument
"""

from __future__ import annotations

import re
from pathlib import Path
from docx import Document

from ir_schema import (
    ParsedDocument, DocumentMeta, DocType, Section,
)

from .docx_profile_detector import detect_profile
from .docx_element_indexer import index_body_elements
from .docx_section_parser import parse_sections


def parse_docx(file_path: str) -> ParsedDocument:
    """
    Parse a DOCX financial statement file into a ParsedDocument IR.

    Steps:
        1. Open the DOCX with python-docx
        2. Auto-detect the DocxProfile (styles, spacing, merge strategy)
        3. Index all body elements (paragraph/table positions)
        4. Split into sections by ABCTitle boundaries
        5. Parse tables within each section
        6. Extract document metadata (company name, periods, doc type)
    """
    doc = Document(file_path)

    # 1. Detect profile
    profile = detect_profile(doc)

    # 2. Index body elements
    body_elements = index_body_elements(doc)

    # 3. Parse sections (includes table parsing)
    sections = parse_sections(doc, body_elements, profile)

    # Update profile section count to actual note count
    total_notes = sum(len(s.notes) for s in sections)
    profile.section_count = total_notes

    # 4. Extract metadata
    meta = _extract_metadata(doc, sections, file_path)

    return ParsedDocument(
        meta=meta,
        sections=sections,
        docx_profile=profile,
    )


def _extract_metadata(
    doc: Document,
    sections: list[Section],
    file_path: str,
) -> DocumentMeta:
    """Extract company name, periods, and doc type from the document."""
    meta = DocumentMeta(source_format="docx")

    # Try to find company name from first few paragraphs
    for para in doc.paragraphs[:20]:
        text = para.text.strip()
        if not text:
            continue

        # Company name heuristics: look for "Co., Ltd" etc. and extract
        # just the company name part (up to the legal suffix)
        lower = text.lower()
        for suffix in ("co., ltd.", "co., ltd", "inc.", "corp.", "corporation"):
            idx = lower.find(suffix)
            if idx >= 0:
                meta.company = text[:idx + len(suffix)].strip().rstrip(",")
                break
        if meta.company:
            break

        # SBL pattern: first non-empty Normal paragraph is company name
        if para.style.name in ("Normal",) and text and not meta.company:
            if len(text) < 80 and not text.startswith("("):
                meta.company = text

    # If no company found, try filename
    if not meta.company:
        fname = Path(file_path).stem
        meta.company = fname.split("_")[0].split(" ")[0]

    # Detect doc type from content
    all_text = " ".join(p.text.lower() for p in doc.paragraphs[:50])
    if "consolidated" in all_text:
        meta.doc_type = DocType.CONSOLIDATED
    elif "separate" in all_text:
        meta.doc_type = DocType.SEPARATE
    else:
        # Check filename
        fname_lower = Path(file_path).name.lower()
        if "별도" in fname_lower or "separate" in fname_lower:
            meta.doc_type = DocType.SEPARATE
        elif "연결" in fname_lower or "consolidat" in fname_lower:
            meta.doc_type = DocType.CONSOLIDATED

    # Detect periods from text
    # Look for "December 31, YYYY and YYYY" or table headers with years
    period_pattern = re.compile(
        r"December\s+31,\s+(\d{4})\s+and\s+(\d{4})", re.IGNORECASE
    )
    for para in doc.paragraphs[:50]:
        text = para.text.strip()
        m = period_pattern.search(text)
        if m:
            y1, y2 = int(m.group(1)), int(m.group(2))
            meta.period_current = str(max(y1, y2))
            meta.period_prior = str(min(y1, y2))
            break

    # Fallback: scan for any two years close together
    if not meta.period_current:
        year_re = re.compile(r"\b(20[12]\d)\b")
        all_years: list[int] = []
        for para in doc.paragraphs[:50]:
            all_years.extend(int(y) for y in year_re.findall(para.text))
        if all_years:
            unique = sorted(set(all_years), reverse=True)
            meta.period_current = str(unique[0])
            if len(unique) >= 2:
                meta.period_prior = str(unique[1])

    return meta
