"""
Auto-detect the formatting profile of a DOCX financial statement.

Examines paragraph styles, table grid columns, merge patterns, and
spacing conventions to build a DocxProfile that describes the document's
formatting strategy.
"""

from __future__ import annotations

from collections import Counter
from docx import Document

from ir_schema import (
    DocxProfile, SpacingStrategy, MergeStrategy, WidthStrategy,
)
from utils.xml_helpers import OOXML_NS, w, find_w, findall_w


# Spacer column width threshold in dxa (twentieths of a point).
# Columns narrower than this are considered spacers.
SPACER_WIDTH_THRESHOLD = 200


def detect_profile(doc: Document) -> DocxProfile:
    """Analyse a python-docx Document and return a DocxProfile."""
    profile = DocxProfile()

    _detect_styles(doc, profile)
    _detect_table_patterns(doc, profile)

    return profile


# ── Style detection ─────────────────────────────────────────────

def _detect_styles(doc: Document, profile: DocxProfile) -> None:
    """Count paragraph styles and assign title / subtitle / body style."""
    style_counter: Counter[str] = Counter()
    for para in doc.paragraphs:
        style_counter[para.style.name] += 1

    if not style_counter:
        return

    profile.primary_style = style_counter.most_common(1)[0][0]

    # Title style: look for ABCTitle (HYBE/SBL share this convention)
    for name in ("ABCTitle",):
        if name in style_counter:
            profile.title_style = name
            break

    # Subtitle style
    for name in ("Subtitle", "aff5"):
        if name in style_counter:
            profile.subtitle_style = name
            break

    # Body style: most common among likely body styles
    body_candidates = [
        s for s in style_counter
        if s not in (profile.title_style, profile.subtitle_style,
                     "_Index", "_Index - Page(s)")
    ]
    if body_candidates:
        profile.body_style = max(body_candidates, key=lambda s: style_counter[s])


# ── Table pattern detection ─────────────────────────────────────

def _detect_table_patterns(doc: Document, profile: DocxProfile) -> None:
    """Analyse all tables for spacing, merge, and width strategies."""
    total_vmerge = 0
    total_gridspan = 0
    tables_with_spacers = 0
    tables_with_explicit_widths = 0
    tables_with_empty_row2 = 0
    spacer_widths_seen: list[int] = []
    total_tables = len(doc.tables)

    if total_tables == 0:
        return

    for tbl_obj in doc.tables:
        tbl = tbl_obj._tbl

        # ── Grid columns ────────────────────────────────
        grid_cols = findall_w(tbl, "w:tblGrid/w:gridCol")
        widths = []
        for gc in grid_cols:
            w_val = gc.get(w("w"))
            if w_val is not None:
                try:
                    widths.append(int(w_val))
                except ValueError:
                    pass

        spacer_cols_in_table = [
            ww for ww in widths if ww < SPACER_WIDTH_THRESHOLD
        ]
        if spacer_cols_in_table:
            tables_with_spacers += 1
            spacer_widths_seen.extend(spacer_cols_in_table)

        if widths:
            tables_with_explicit_widths += 1

        # ── Merge counts ────────────────────────────────
        vmerge_els = tbl.findall(f".//{w('vMerge')}")
        gridspan_els = tbl.findall(f".//{w('gridSpan')}")
        total_vmerge += len(vmerge_els)
        total_gridspan += len(gridspan_els)

        # ── Empty row-2 pattern ─────────────────────────
        rows = findall_w(tbl, "w:tr")
        if len(rows) >= 3:
            row2 = rows[1]
            cells = findall_w(row2, "w:tc")
            all_empty = all(
                _cell_text(c).strip() == "" for c in cells
            )
            if all_empty:
                tables_with_empty_row2 += 1

    # ── Assign strategies ───────────────────────────────
    spacer_ratio = tables_with_spacers / total_tables
    empty_row_ratio = tables_with_empty_row2 / total_tables

    if spacer_ratio > 0.3:
        profile.spacing_strategy = SpacingStrategy.SPACER_COLUMN
    elif empty_row_ratio > 0.2:
        profile.spacing_strategy = SpacingStrategy.EMPTY_ROW
    elif spacer_ratio > 0.1 and empty_row_ratio > 0.1:
        profile.spacing_strategy = SpacingStrategy.MIXED
    else:
        profile.spacing_strategy = SpacingStrategy.NONE

    profile.empty_row_pattern = empty_row_ratio > 0.2
    profile.spacer_col_widths = sorted(set(spacer_widths_seen))

    # Merge strategy
    total_merge = total_vmerge + total_gridspan
    if total_merge == 0:
        profile.merge_strategy = MergeStrategy.MINIMAL
    elif total_vmerge > total_gridspan * 2:
        profile.merge_strategy = MergeStrategy.VMERGE_HEAVY
    elif total_gridspan > total_vmerge * 2:
        profile.merge_strategy = MergeStrategy.GRIDSPAN_HEAVY
    else:
        profile.merge_strategy = MergeStrategy.BALANCED

    # Width strategy
    width_ratio = tables_with_explicit_widths / total_tables
    if width_ratio > 0.8:
        profile.width_strategy = WidthStrategy.FIXED
    elif width_ratio < 0.2:
        profile.width_strategy = WidthStrategy.AUTO
    else:
        profile.width_strategy = WidthStrategy.MIXED

    profile.section_count = total_tables  # will be refined by section parser


def _cell_text(tc_element) -> str:
    """Extract text from a w:tc element."""
    texts = []
    for t_el in tc_element.iter(w("t")):
        if t_el.text:
            texts.append(t_el.text)
    return "".join(texts)
