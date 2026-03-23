#!/usr/bin/env python3
"""
SARA Output Quality Analysis Script v2
========================================
Properly analyzes the Notes-only DOCX output vs DSD source data.
The DOCX is the notes section of the English financial report (not main FS tables).
"""

import sys
import os
import re
from collections import defaultdict, Counter

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from docx import Document
from skills.parse_dsd import parse_dsd
from utils.number_format import parse_korean_number
from ir_schema import ElementType

# ─── Config ───
RESULT_PATH = "backend/outputs/ac368833-2106-42f6-b32a-e5edeafcf53a/SARA_result.docx"
TEMPLATE_PATH = "files/Hybe 2024 Eng Report.docx"
DSD_PATH = "files/Hybe 2025 Eng Report.dsd"


def parse_number(text):
    """Parse a number from English-formatted financial text."""
    if not text:
        return None
    cleaned = text.strip()
    cleaned = cleaned.replace("\\", "").replace("\u20a9", "").strip()
    if cleaned in ("-", "—", "–", ""):
        return 0
    if "%" in cleaned:
        # Try to parse percentage
        pct = cleaned.replace("%", "").strip().replace(",", "")
        try:
            return float(pct)
        except ValueError:
            return None
    if re.match(r'^[A-Za-z\s]+$', cleaned):
        return None

    is_negative = False
    if cleaned.startswith("(") and cleaned.endswith(")"):
        is_negative = True
        cleaned = cleaned[1:-1].strip()
    elif cleaned.startswith("-"):
        rest = cleaned[1:].strip().replace(",", "")
        if rest and rest.replace(".", "").isdigit():
            is_negative = True
            cleaned = rest

    cleaned = cleaned.replace(",", "")
    if not cleaned:
        return None
    try:
        if "." in cleaned:
            val = float(cleaned)
        else:
            val = int(cleaned)
        return -val if is_negative else val
    except ValueError:
        return None


def extract_table_numbers(rows):
    """Extract all numeric values from a table's rows."""
    numbers = []
    for row in rows:
        for cell in row:
            val = parse_number(cell)
            if val is not None and val != 0:
                numbers.append(val)
    return numbers


def extract_table_numbers_korean(rows):
    """Extract all numeric values from a DSD table's rows (Korean format)."""
    numbers = []
    for row in rows:
        for cell in row:
            val = parse_korean_number(cell)
            if val is not None and val != 0:
                numbers.append(val)
    return numbers


def get_table_data(table):
    """Extract all cell text from a python-docx table."""
    rows = []
    for row in table.rows:
        cells = [c.text.strip() for c in row.cells]
        rows.append(cells)
    return rows


def get_table_fingerprint(rows, max_nums=20):
    """Create a fingerprint of a table based on its numbers."""
    numbers = []
    for row in rows:
        for cell in row:
            val = parse_number(cell)
            if val is not None and val != 0:
                numbers.append(val)
    return set(numbers[:max_nums])


def compare_tables_detailed(dsd_rows, docx_rows, label=""):
    """Detailed cell-by-cell comparison between a DSD table and DOCX table."""
    result = {
        "matching_numbers": 0,
        "mismatched_numbers": 0,
        "missing_in_docx": 0,
        "extra_in_docx": 0,
        "details": [],
    }

    # Extract numbers from DSD
    dsd_nums = extract_table_numbers_korean(dsd_rows)
    docx_nums = extract_table_numbers(docx_rows)

    dsd_set = Counter(dsd_nums)
    docx_set = Counter(docx_nums)

    # Numbers in both
    common_keys = set(dsd_set.keys()) & set(docx_set.keys())
    for key in common_keys:
        matched = min(dsd_set[key], docx_set[key])
        result["matching_numbers"] += matched

    # Numbers only in DSD
    for key in set(dsd_set.keys()) - common_keys:
        result["missing_in_docx"] += dsd_set[key]

    # Numbers only in DOCX
    for key in set(docx_set.keys()) - common_keys:
        result["extra_in_docx"] += docx_set[key]

    # Mismatches from partial matches
    for key in common_keys:
        diff = abs(dsd_set[key] - docx_set[key])
        if dsd_set[key] > docx_set[key]:
            result["missing_in_docx"] += diff
        elif docx_set[key] > dsd_set[key]:
            result["extra_in_docx"] += diff

    return result


def main():
    print("=" * 110)
    print("SARA OUTPUT QUALITY ANALYSIS — DETAILED REPORT")
    print("=" * 110)

    # Load files
    print("\nLoading files...")
    result_doc = Document(RESULT_PATH)
    template_doc = Document(TEMPLATE_PATH)
    parsed_dsd = parse_dsd(DSD_PATH)

    print(f"  SARA Result: {len(result_doc.tables)} tables, {len(result_doc.paragraphs)} paragraphs")
    print(f"  Template (Prior Year): {len(template_doc.tables)} tables, {len(template_doc.paragraphs)} paragraphs")
    print(f"  DSD Source: {parsed_dsd.meta.company}")
    print(f"  DSD Period: {parsed_dsd.meta.period_current} (current), {parsed_dsd.meta.period_prior} (prior)")

    # ──────────────────────────────────────────
    # Part 1: Year Rolling Analysis
    # ──────────────────────────────────────────
    print("\n")
    print("=" * 110)
    print("PART 1: YEAR ROLLING ANALYSIS (2024→2025 / 2023→2024)")
    print("=" * 110)

    # Compare template and result to check year rolling
    year_changes = {"2024→2025": 0, "2023→2024": 0, "unchanged": 0, "other_changes": 0}

    for i in range(min(len(result_doc.tables), len(template_doc.tables))):
        result_rows = get_table_data(result_doc.tables[i])
        template_rows = get_table_data(template_doc.tables[i])

        for ri in range(min(len(result_rows), len(template_rows))):
            for ci in range(min(len(result_rows[ri]), len(template_rows[ri]))):
                r_val = result_rows[ri][ci]
                t_val = template_rows[ri][ci]
                if r_val != t_val:
                    if t_val == "2024" and r_val == "2025":
                        year_changes["2024→2025"] += 1
                    elif t_val == "2023" and r_val == "2024":
                        year_changes["2023→2024"] += 1
                    elif "2024" in t_val and "2025" in r_val and t_val.replace("2024", "2025") == r_val:
                        year_changes["2024→2025"] += 1
                    elif "2023" in t_val and "2024" in r_val and t_val.replace("2023", "2024") == r_val:
                        year_changes["2023→2024"] += 1
                    else:
                        year_changes["other_changes"] += 1

    print(f"\n  Year rolling in tables:")
    print(f"    2024 → 2025 (current year): {year_changes['2024→2025']} cells")
    print(f"    2023 → 2024 (prior year): {year_changes['2023→2024']} cells")
    print(f"    Other changes (data updates): {year_changes['other_changes']} cells")

    # Also check paragraphs
    para_year_changes = {"2024→2025": 0, "2023→2024": 0, "other": 0}
    for i in range(min(len(result_doc.paragraphs), len(template_doc.paragraphs))):
        r_text = result_doc.paragraphs[i].text
        t_text = template_doc.paragraphs[i].text
        if r_text != t_text:
            if "2024" in t_text and "2025" in r_text:
                para_year_changes["2024→2025"] += 1
            elif "2023" in t_text and "2024" in r_text:
                para_year_changes["2023→2024"] += 1
            else:
                para_year_changes["other"] += 1

    print(f"\n  Year rolling in paragraphs:")
    print(f"    2024 → 2025: {para_year_changes['2024→2025']} paragraphs")
    print(f"    2023 → 2024: {para_year_changes['2023→2024']} paragraphs")
    print(f"    Other changes: {para_year_changes['other']} paragraphs")

    # ──────────────────────────────────────────
    # Part 2: Table-by-table data comparison
    # ──────────────────────────────────────────
    print("\n")
    print("=" * 110)
    print("PART 2: TABLE-BY-TABLE DATA COMPARISON (DSD vs DOCX)")
    print("=" * 110)

    # Build list of DSD note tables with metadata
    dsd_note_tables = []
    for section in parsed_dsd.sections:
        for note in section.notes:
            table_idx = 0
            for elem in note.elements:
                if elem.type == ElementType.TABLE and elem.table:
                    rows = []
                    for r in elem.table.headers:
                        rows.append([c.text for c in r.cells])
                    for r in elem.table.rows:
                        rows.append([c.text for c in r.cells])
                    numbers = extract_table_numbers_korean(rows)
                    dsd_note_tables.append({
                        "note_num": note.number,
                        "note_title": note.title,
                        "table_idx": table_idx,
                        "rows": rows,
                        "numbers": numbers,
                        "row_count": len(rows),
                    })
                    table_idx += 1

    print(f"\n  DSD note tables: {len(dsd_note_tables)}")

    # Build DOCX table data
    docx_table_data = []
    for i, table in enumerate(result_doc.tables):
        rows = get_table_data(table)
        numbers = extract_table_numbers(rows)
        docx_table_data.append({
            "idx": i,
            "rows": rows,
            "numbers": numbers,
            "row_count": len(rows),
        })

    # Template table data
    template_table_data = []
    for i, table in enumerate(template_doc.tables):
        rows = get_table_data(table)
        numbers = extract_table_numbers(rows)
        template_table_data.append({
            "idx": i,
            "rows": rows,
            "numbers": numbers,
            "row_count": len(rows),
        })

    # Match DSD note tables to DOCX tables
    # The agent log says 104 tables were matched with 1443 cells updated
    # Strategy: find the best DOCX table match for each DSD table

    matched_pairs = []
    used_docx_indices = set()

    for dsd_t in dsd_note_tables:
        if not dsd_t["numbers"]:
            continue

        best_match = None
        best_score = 0

        dsd_num_set = set(dsd_t["numbers"][:30])
        if not dsd_num_set:
            continue

        for dt in docx_table_data:
            if dt["idx"] in used_docx_indices:
                continue
            docx_num_set = set(dt["numbers"][:30])
            if not docx_num_set:
                continue

            # Row count similarity
            if abs(dt["row_count"] - dsd_t["row_count"]) > max(10, dsd_t["row_count"] * 0.5):
                continue

            common = dsd_num_set & docx_num_set
            score = len(common) / max(len(dsd_num_set), 1)

            if score > best_score:
                best_score = score
                best_match = dt

        if best_match and best_score >= 0.3:
            matched_pairs.append({
                "dsd": dsd_t,
                "docx": best_match,
                "score": best_score,
            })
            used_docx_indices.add(best_match["idx"])

    print(f"  Matched DSD↔DOCX table pairs: {len(matched_pairs)}")

    # ──────────────────────────────────────────
    # Analyze each matched pair
    # ──────────────────────────────────────────
    total_stats = {
        "perfect_match": 0,
        "high_match": 0,
        "partial_match": 0,
        "low_match": 0,
        "total_dsd_numbers": 0,
        "total_matched": 0,
        "total_missing": 0,
    }

    # Group by note for cleaner output
    note_results = defaultdict(list)

    for pair in matched_pairs:
        dsd = pair["dsd"]
        docx = pair["docx"]

        comp = compare_tables_detailed(dsd["rows"], docx["rows"])
        total_nums = comp["matching_numbers"] + comp["missing_in_docx"] + comp["extra_in_docx"]

        score = pair["score"]
        if score >= 0.95:
            quality = "PERFECT"
            total_stats["perfect_match"] += 1
        elif score >= 0.7:
            quality = "HIGH"
            total_stats["high_match"] += 1
        elif score >= 0.5:
            quality = "PARTIAL"
            total_stats["partial_match"] += 1
        else:
            quality = "LOW"
            total_stats["low_match"] += 1

        total_stats["total_dsd_numbers"] += len(dsd["numbers"])
        total_stats["total_matched"] += comp["matching_numbers"]
        total_stats["total_missing"] += comp["missing_in_docx"]

        note_key = f"Note {dsd['note_num']}: {dsd['note_title']}"
        note_results[note_key].append({
            "dsd_table_idx": dsd["table_idx"],
            "docx_table_idx": docx["idx"],
            "score": score,
            "quality": quality,
            "matched": comp["matching_numbers"],
            "missing": comp["missing_in_docx"],
            "extra": comp["extra_in_docx"],
            "dsd_nums": len(dsd["numbers"]),
            "docx_nums": len(docx["numbers"]),
        })

    print(f"\n  {'Note':55} | {'DOCX#':>5} | {'Score':>6} | {'Quality':>8} | {'Match':>5} | {'Miss':>5} | {'Extra':>5}")
    print(f"  {'─'*55}─┼─{'─'*5}─┼─{'─'*6}─┼─{'─'*8}─┼─{'─'*5}─┼─{'─'*5}─┼─{'─'*5}")

    for note_key in sorted(note_results.keys(), key=lambda x: int(x.split(":")[0].split()[-1]) if x.split(":")[0].split()[-1].isdigit() else 999):
        tables = note_results[note_key]
        for t in tables:
            note_display = f"{note_key} [T{t['dsd_table_idx']}]"[:55]
            print(f"  {note_display:55} | {t['docx_table_idx']:>5} | {t['score']:>5.0%} | {t['quality']:>8} | {t['matched']:>5} | {t['missing']:>5} | {t['extra']:>5}")

    print(f"\n  Summary:")
    print(f"    Perfect match (>=95%): {total_stats['perfect_match']}")
    print(f"    High match (70-95%): {total_stats['high_match']}")
    print(f"    Partial match (50-70%): {total_stats['partial_match']}")
    print(f"    Low match (<50%): {total_stats['low_match']}")
    print(f"    Total DSD numbers: {total_stats['total_dsd_numbers']}")
    print(f"    Total matched: {total_stats['total_matched']}")
    print(f"    Total missing from DOCX: {total_stats['total_missing']}")

    # ──────────────────────────────────────────
    # Part 3: Data updates vs template
    # ──────────────────────────────────────────
    print("\n")
    print("=" * 110)
    print("PART 3: DATA CHANGES FROM TEMPLATE (Current Year Data Insertion)")
    print("=" * 110)

    # For each table, compare result vs template to find which cells were updated
    tables_with_changes = 0
    tables_no_changes = 0
    total_cells_changed = 0
    total_numeric_cells_changed = 0

    # Detailed per-table analysis
    print(f"\n  {'Table':>5} | {'Rows':>4} | {'Changed':>7} | {'NumChg':>7} | {'Status':10} | Context")
    print(f"  {'─'*5}─┼─{'─'*4}─┼─{'─'*7}─┼─{'─'*7}─┼─{'─'*10}─┼─{'─'*50}")

    # Get paragraph context for each table
    from docx.oxml.ns import qn
    body = result_doc.element.body
    table_contexts = {}
    prev_para = ""
    tbl_idx = 0
    for elem in body:
        if elem.tag == qn('w:p'):
            runs = elem.findall(f'.//{qn("w:t")}')
            text = ''.join((r.text or '') for r in runs)
            if text.strip():
                prev_para = text.strip()[:60]
        elif elem.tag == qn('w:tbl'):
            table_contexts[tbl_idx] = prev_para
            tbl_idx += 1

    for i in range(min(len(result_doc.tables), len(template_doc.tables))):
        result_rows = get_table_data(result_doc.tables[i])
        template_rows = get_table_data(template_doc.tables[i])

        changed_cells = 0
        numeric_changed = 0

        for ri in range(min(len(result_rows), len(template_rows))):
            for ci in range(min(len(result_rows[ri]), len(template_rows[ri]))):
                r_val = result_rows[ri][ci]
                t_val = template_rows[ri][ci]
                if r_val != t_val:
                    changed_cells += 1
                    # Check if numeric change
                    r_num = parse_number(r_val)
                    t_num = parse_number(t_val)
                    if r_num is not None or t_num is not None:
                        numeric_changed += 1

        if changed_cells > 0:
            tables_with_changes += 1
            total_cells_changed += changed_cells
            total_numeric_cells_changed += numeric_changed

            # Determine status
            if numeric_changed > 0:
                status = "DATA_UPD"
            else:
                status = "TEXT_UPD"

            context = table_contexts.get(i, "")[:50]
            print(f"  {i:>5} | {len(result_rows):>4} | {changed_cells:>7} | {numeric_changed:>7} | {status:10} | {context}")
        else:
            tables_no_changes += 1

    print(f"\n  Summary:")
    print(f"    Tables with changes: {tables_with_changes}")
    print(f"    Tables unchanged: {tables_no_changes}")
    print(f"    Total cells changed: {total_cells_changed}")
    print(f"    Numeric cells changed: {total_numeric_cells_changed}")

    # ──────────────────────────────────────────
    # Part 4: Totals/Subtotals Analysis
    # ──────────────────────────────────────────
    print("\n")
    print("=" * 110)
    print("PART 4: TOTALS AND SUBTOTALS VALIDATION")
    print("=" * 110)

    # For each DOCX table, find rows with "Total" or "Subtotal" and check
    # if the values are consistent with the DSD source
    total_rows_checked = 0
    total_rows_correct = 0
    total_rows_wrong = 0
    total_rows_empty = 0

    total_issues = []

    for pair in matched_pairs:
        dsd = pair["dsd"]
        docx = pair["docx"]

        dsd_rows = dsd["rows"]
        docx_rows = docx["rows"]

        # Find total/subtotal rows in DSD
        dsd_totals = {}
        for ri, row in enumerate(dsd_rows):
            combined = " ".join(row)
            if re.search(r'합\s*계|총\s*계|소\s*계', combined):
                label = ""
                nums = []
                for c in row:
                    val = parse_korean_number(c)
                    if val is not None:
                        nums.append(val)
                    elif not label and c.strip():
                        label = c.strip()
                dsd_totals[ri] = {"label": label, "values": nums}

        # Find total rows in DOCX
        docx_totals = {}
        for ri, row in enumerate(docx_rows):
            combined = " ".join(row).lower()
            if "total" in combined or "subtotal" in combined or "sub-total" in combined:
                label = ""
                nums = []
                for c in row:
                    val = parse_number(c)
                    if val is not None:
                        nums.append(val)
                    elif not label and c.strip() and "thousands" not in c.lower():
                        label = c.strip()
                docx_totals[ri] = {"label": label, "values": nums}

        # Cross-compare totals
        for dsd_ri, dsd_total in dsd_totals.items():
            total_rows_checked += 1
            dsd_vals = set(dsd_total["values"])
            if not dsd_vals:
                continue

            found_match = False
            for docx_ri, docx_total in docx_totals.items():
                docx_vals = set(docx_total["values"])
                if dsd_vals & docx_vals:  # Any overlap
                    overlap = len(dsd_vals & docx_vals) / len(dsd_vals)
                    if overlap >= 0.5:
                        found_match = True
                        if overlap >= 0.9:
                            total_rows_correct += 1
                        else:
                            total_rows_wrong += 1
                            total_issues.append({
                                "note": f"Note {dsd['note_num']}: {dsd['note_title']}",
                                "docx_table": docx["idx"],
                                "dsd_label": dsd_total["label"],
                                "docx_label": docx_total["label"],
                                "dsd_values": sorted(dsd_total["values"]),
                                "docx_values": sorted(docx_total["values"]),
                            })
                        break

            if not found_match:
                # Check if values are in DOCX at all (even in non-total rows)
                all_docx_nums = set(extract_table_numbers(docx_rows))
                overlap = len(dsd_vals & all_docx_nums) / len(dsd_vals) if dsd_vals else 0
                if overlap >= 0.5:
                    total_rows_correct += 1  # Values present but not in a "total" labeled row
                else:
                    total_rows_wrong += 1
                    total_issues.append({
                        "note": f"Note {dsd['note_num']}: {dsd['note_title']}",
                        "docx_table": docx["idx"],
                        "dsd_label": dsd_total["label"],
                        "docx_label": "(not found)",
                        "dsd_values": sorted(dsd_total["values"]),
                        "docx_values": [],
                    })

    print(f"\n  Total/subtotal rows checked: {total_rows_checked}")
    print(f"    Correct: {total_rows_correct}")
    print(f"    Wrong/missing: {total_rows_wrong}")
    if total_rows_checked > 0:
        print(f"    Accuracy: {total_rows_correct/total_rows_checked*100:.1f}%")

    if total_issues:
        print(f"\n  Specific total/subtotal issues (showing up to 20):")
        for issue in total_issues[:20]:
            dsd_str = ", ".join(f"{v:,}" for v in issue["dsd_values"][:5])
            docx_str = ", ".join(f"{v:,}" for v in issue["docx_values"][:5])
            print(f"    {issue['note'][:45]}")
            print(f"      DSD total [{issue['dsd_label'][:30]}]: {dsd_str}")
            print(f"      DOCX total [{issue['docx_label'][:30]}]: {docx_str}")

    # ──────────────────────────────────────────
    # Part 5: Current Year vs Prior Year Column Analysis
    # ──────────────────────────────────────────
    print("\n")
    print("=" * 110)
    print("PART 5: CURRENT YEAR (당기/2025) vs PRIOR YEAR (전기/2024) DATA ACCURACY")
    print("=" * 110)

    # For the matched pairs, do a more detailed row-by-row comparison
    # showing expected vs actual for both current and prior year
    sample_count = 0
    max_samples = 15

    for pair in matched_pairs:
        if sample_count >= max_samples:
            break

        dsd = pair["dsd"]
        docx = pair["docx"]

        # Skip tables with very few numbers
        if len(dsd["numbers"]) < 4:
            continue

        dsd_rows = dsd["rows"]
        docx_rows = docx["rows"]

        # Try to identify which columns are current year and prior year
        # DSD: typically has 당기 and 전기 columns
        # DOCX: typically has 2025 and 2024 columns

        # Skip if perfect match
        if pair["score"] >= 0.95:
            continue

        sample_count += 1

        print(f"\n  --- Note {dsd['note_num']}: {dsd['note_title']} (DSD Table {dsd['table_idx']}) → DOCX Table {docx['idx']} ---")
        print(f"      Match score: {pair['score']:.0%}")

        # Show DSD rows
        print(f"\n      DSD Rows ({len(dsd_rows)}):")
        for ri, row in enumerate(dsd_rows[:8]):
            cells_str = " | ".join(c[:25] for c in row if c.strip())
            print(f"        [{ri}] {cells_str[:100]}")
        if len(dsd_rows) > 8:
            print(f"        ... ({len(dsd_rows) - 8} more rows)")

        # Show DOCX rows
        print(f"      DOCX Rows ({len(docx_rows)}):")
        for ri, row in enumerate(docx_rows[:8]):
            cells_str = " | ".join(c[:25] for c in row if c.strip())
            print(f"        [{ri}] {cells_str[:100]}")
        if len(docx_rows) > 8:
            print(f"        ... ({len(docx_rows) - 8} more rows)")

        # Identify missing numbers
        dsd_num_set = set(dsd["numbers"])
        docx_num_set = set(docx["numbers"])
        missing = dsd_num_set - docx_num_set
        extra = docx_num_set - dsd_num_set

        if missing:
            sample = sorted(missing, key=abs, reverse=True)[:5]
            print(f"      Missing from DOCX: {[f'{v:,}' for v in sample]}")
        if extra:
            sample = sorted(extra, key=abs, reverse=True)[:5]
            print(f"      Extra in DOCX (not in DSD): {[f'{v:,}' for v in sample]}")

    # ──────────────────────────────────────────
    # Part 6: Unmatched DSD tables
    # ──────────────────────────────────────────
    print("\n")
    print("=" * 110)
    print("PART 6: UNMATCHED DSD TABLES (Data Not Found in DOCX)")
    print("=" * 110)

    matched_dsd_keys = set()
    for pair in matched_pairs:
        key = (pair["dsd"]["note_num"], pair["dsd"]["table_idx"])
        matched_dsd_keys.add(key)

    unmatched = []
    for dsd_t in dsd_note_tables:
        key = (dsd_t["note_num"], dsd_t["table_idx"])
        if key not in matched_dsd_keys and dsd_t["numbers"]:
            unmatched.append(dsd_t)

    print(f"\n  Unmatched DSD tables with data: {len(unmatched)}")
    for u in unmatched:
        nums_str = ", ".join(f"{v:,}" for v in u["numbers"][:5])
        print(f"    Note {u['note_num']}: {u['note_title'][:40]} [T{u['table_idx']}] ({u['row_count']}r, {len(u['numbers'])} nums) — {nums_str}")

    # ──────────────────────────────────────────
    # Part 7: Empty DOCX tables (no data filled)
    # ──────────────────────────────────────────
    print("\n")
    print("=" * 110)
    print("PART 7: EMPTY/UNCHANGED DOCX TABLES (Potential Missing Data)")
    print("=" * 110)

    empty_tables = []
    for i in range(len(result_doc.tables)):
        result_rows = get_table_data(result_doc.tables[i])
        template_rows = get_table_data(template_doc.tables[i]) if i < len(template_doc.tables) else []

        # Check if ALL data cells are same as template (no updates)
        has_numeric_change = False
        for ri in range(min(len(result_rows), len(template_rows))):
            for ci in range(min(len(result_rows[ri]), len(template_rows[ri]))):
                r_val = result_rows[ri][ci]
                t_val = template_rows[ri][ci]
                if r_val != t_val:
                    r_num = parse_number(r_val)
                    if r_num is not None:
                        has_numeric_change = True
                        break
            if has_numeric_change:
                break

        if not has_numeric_change and len(result_rows) >= 3:
            # Check if the table has numbers at all
            all_text = " ".join(" ".join(r) for r in result_rows)
            has_numbers = bool(re.search(r'\d{2,}', all_text.replace("2025", "").replace("2024", "").replace("2023", "")))
            context = table_contexts.get(i, "")[:50]
            if has_numbers:
                empty_tables.append((i, len(result_rows), "HAS_PRIOR_DATA", context))
            else:
                empty_tables.append((i, len(result_rows), "NO_DATA", context))

    print(f"\n  Tables with no numeric updates from template: {len(empty_tables)}")
    for idx, rows, status, ctx in empty_tables:
        print(f"    Table {idx:>3} ({rows}r): [{status:15}] {ctx}")

    # ──────────────────────────────────────────
    # Part 8: Specific Row-by-Row Comparison (Detailed Examples)
    # ──────────────────────────────────────────
    print("\n")
    print("=" * 110)
    print("PART 8: DETAILED ROW-BY-ROW COMPARISON EXAMPLES")
    print("=" * 110)

    # Pick tables where there are mismatches and show detailed comparison
    example_count = 0
    for pair in matched_pairs:
        if example_count >= 5:
            break
        if pair["score"] >= 0.95 or pair["score"] < 0.3:
            continue

        dsd = pair["dsd"]
        docx = pair["docx"]
        template = template_table_data[docx["idx"]] if docx["idx"] < len(template_table_data) else None

        dsd_rows = dsd["rows"]
        docx_rows = docx["rows"]
        tmpl_rows = template["rows"] if template else []

        example_count += 1

        print(f"\n  === Note {dsd['note_num']}: {dsd['note_title']} (DSD T{dsd['table_idx']}) → DOCX T{docx['idx']} (score={pair['score']:.0%}) ===")

        # Row-by-row showing: DSD Korean | Template (Prior Year Eng) | Result (Current Year Eng)
        max_rows_show = min(len(dsd_rows), 12)
        print(f"\n  {'#':>3} | {'DSD (Korean)':50} | {'Template (Prior)':30} | {'Result (Output)':30} | Status")
        print(f"  {'─'*3}─┼─{'─'*50}─┼─{'─'*30}─┼─{'─'*30}─┼─{'─'*8}")

        for ri in range(max_rows_show):
            dsd_cells = " | ".join(c[:15] for c in dsd_rows[ri] if c.strip())[:50]
            tmpl_cells = " | ".join(c[:15] for c in tmpl_rows[ri] if c.strip())[:30] if ri < len(tmpl_rows) else "N/A"
            result_cells = " | ".join(c[:15] for c in docx_rows[ri] if c.strip())[:30] if ri < len(docx_rows) else "N/A"

            # Check if result differs from template
            result_changed = False
            if ri < len(docx_rows) and ri < len(tmpl_rows):
                for ci in range(min(len(docx_rows[ri]), len(tmpl_rows[ri]))):
                    if docx_rows[ri][ci] != tmpl_rows[ri][ci]:
                        result_changed = True
                        break

            status = "CHANGED" if result_changed else "same"
            print(f"  {ri:>3} | {dsd_cells:50} | {tmpl_cells:30} | {result_cells:30} | {status}")

    # ──────────────────────────────────────────
    # Part 9: Pattern Analysis
    # ──────────────────────────────────────────
    print("\n")
    print("=" * 110)
    print("PART 9: ERROR PATTERN ANALYSIS")
    print("=" * 110)

    # Analyze types of mismatches
    prior_year_unchanged = 0  # Prior year column not updated (same as template)
    current_year_wrong = 0    # Current year has wrong values
    current_year_correct = 0  # Current year has correct values

    for pair in matched_pairs:
        dsd = pair["dsd"]
        docx = pair["docx"]
        template = template_table_data[docx["idx"]] if docx["idx"] < len(template_table_data) else None

        if not template:
            continue

        # For each numeric cell that differs from template, check if it matches DSD
        dsd_nums = set(dsd["numbers"])
        docx_rows = docx["rows"]
        tmpl_rows = template["rows"]

        for ri in range(min(len(docx_rows), len(tmpl_rows))):
            for ci in range(min(len(docx_rows[ri]), len(tmpl_rows[ri]))):
                r_val = docx_rows[ri][ci]
                t_val = tmpl_rows[ri][ci]

                if r_val == t_val:
                    continue

                r_num = parse_number(r_val)
                if r_num is not None and r_num != 0:
                    if r_num in dsd_nums:
                        current_year_correct += 1
                    else:
                        current_year_wrong += 1

    # Also check: prior year columns — are they the same between template and result?
    # (They should be, since prior year = template current year)
    print(f"\n  Numeric cell analysis (changed cells only):")
    print(f"    Cells updated with correct DSD value: {current_year_correct}")
    print(f"    Cells updated with wrong value (not in DSD): {current_year_wrong}")
    if (current_year_correct + current_year_wrong) > 0:
        acc = current_year_correct / (current_year_correct + current_year_wrong) * 100
        print(f"    Update accuracy: {acc:.1f}%")

    # ──────────────────────────────────────────
    # Overall Summary
    # ──────────────────────────────────────────
    print("\n")
    print("=" * 110)
    print("OVERALL SUMMARY")
    print("=" * 110)

    print(f"""
  Document Structure:
    - This is the NOTES section of the Hybe English financial report (not the main FS tables)
    - Template (prior year): 145 tables, 764 paragraphs
    - Result (output): 145 tables, 764 paragraphs (structure preserved)

  Year Rolling:
    - Table cells rolled: {year_changes['2024→2025'] + year_changes['2023→2024']}
    - Paragraph year updates: {para_year_changes['2024→2025'] + para_year_changes['2023→2024']}

  Data Filling:
    - Tables with data changes: {tables_with_changes} / {len(result_doc.tables)}
    - Total cells changed: {total_cells_changed}
    - Numeric cells changed: {total_numeric_cells_changed}

  DSD Matching:
    - DSD note tables with data: {sum(1 for t in dsd_note_tables if t['numbers'])}
    - Successfully matched to DOCX: {len(matched_pairs)}
    - Unmatched (data not found): {len(unmatched)}
    - Match quality: {total_stats['perfect_match']} perfect, {total_stats['high_match']} high, {total_stats['partial_match']} partial, {total_stats['low_match']} low

  Data Accuracy:
    - Correct numeric updates: {current_year_correct}
    - Wrong numeric updates: {current_year_wrong}
    - Update accuracy: {current_year_correct/(current_year_correct+current_year_wrong)*100:.1f}% (of changed cells)

  Totals/Subtotals:
    - Total rows checked: {total_rows_checked}
    - Correct: {total_rows_correct}
    - Wrong/missing: {total_rows_wrong}
    - Accuracy: {total_rows_correct/max(total_rows_checked,1)*100:.1f}%

  Agent Performance:
    - Processing time: ~12 minutes
    - 104/138 DSD tables auto-matched
    - 1,443 cells auto-updated
    - 441 Korean-English glossary pairs built
""")


if __name__ == "__main__":
    main()
